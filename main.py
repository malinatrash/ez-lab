from fastapi import FastAPI
from fastapi.responses import FileResponse
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

from starlette.middleware.cors import CORSMiddleware

from models.ReportRequest import ReportRequest

app = FastAPI()

# CORS middleware setup
origins = [
    "http://localhost:5500",
    "http://localhost:3000",
    "http://92.38.241.219:5500",
    "http://92.38.241.219:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/report/")
async def generate_report(request: ReportRequest):
    year = datetime.now().year

    doc = Document('template.docx')

    # Замена заполнителей в абзацах
    for paragraph in doc.paragraphs:
        for placeholder, value in {
            '%институт%': request.institute,
            '%подразделение%': request.podr,
            '%группа%': request.group,
            '%фио%': request.fio,
            '%фиопрепода%': request.fio_prepod,
            '%должность%': request.role,
            '%тип%': request.work_type,
            '%дисциплина%': request.discipline,
            '%тема%': request.topic,
            '%год%': str(year)
        }.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                for run in paragraph.runs:
                    run.font.size = Pt(14)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in {
                        '%институт%': request.institute,
                        '%подразделение%': request.podr,
                        '%группа%': request.group,
                        '%фио%': request.fio,
                        '%фиопрепода%': request.fio_prepod,
                        '%должность%': request.role,
                        '%тип%': request.work_type,
                        '%дисциплина%': request.discipline,
                        '%тема%': request.topic,
                        '%год%': str(year)
                    }.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            for run in paragraph.runs:
                                run.font.size = Pt(14)

    for i, chapter in enumerate(request.chapters):
        doc.add_page_break()

        # Создать новый абзац для названия главы
        chapter_paragraph = doc.add_paragraph(chapter, style="Normal")
        run = chapter_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        chapter_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for additional_chapter in ["Заключение", "Список использованных источников"]:
        doc.add_page_break()
        additional_paragraph = doc.add_paragraph(additional_chapter, style="Normal")
        run = additional_paragraph.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
        additional_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    output_dir = 'reports'
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'{request.discipline}_{request.topic}.docx')
    doc.save(output_path)

    return FileResponse(output_path,
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        filename=f'{request.discipline}_{request.topic}.docx')

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
